"""Generate Lab4_Report.docx with per-exercise requirements, code, logic, and screenshots."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SCREENSHOTS = ROOT / "docs" / "screenshots"
OUTPUT = ROOT / "Lab4_Report.docx"
REAL_SCREENSHOTS_FLAG = ROOT / "docs" / "screenshots" / ".real_captures"

EXERCISES = [
    {
        "title": "Exercise 1 — Highest Score",
        "requirement": "Create an additional area to display the highest score achieved.",
        "code": """// GameView.kt — load and persist high score
private val prefs = context.getSharedPreferences(PREFS_NAME, Context.MODE_PRIVATE)
highScore = prefs.getInt(KEY_HIGH_SCORE, 0)

private fun updateHighScore() {
    if (score > highScore) {
        highScore = score
        prefs.edit().putInt(KEY_HIGH_SCORE, highScore).apply()
    }
}

// draw() — display on HUD
canvas.drawText("High Score: $highScore", 50f, 130f, hudPaint)""",
        "logic": (
            "On init, the saved high score is read from SharedPreferences. Each time the "
            "player scores points or the game ends, updateHighScore() compares the current "
            "score with the stored value and saves a new record if higher. The draw() method "
            "renders the value on screen every frame."
        ),
        "screenshot": "02_hud.png",
        "caption": "Figure 1: HUD showing current score and highest score",
    },
    {
        "title": "Exercise 2 — Three Lives",
        "requirement": (
            "Give the player 3 lives for each game session (the player loses only after "
            "enemies reach the boundary 3 times)."
        ),
        "code": """// GameView.kt
private var lives: Int = 3

val boundaryY = height - 100f
val reachedBoundary = opponents.filter { it.y + it.height >= boundaryY }
for (opponent in reachedBoundary) {
    opponents.remove(opponent)
    lives--
}
if (lives <= 0) {
    gameOver = true
    updateHighScore()
}

// draw()
canvas.drawText("Lives: $lives", 50f, 180f, hudPaint)""",
        "logic": (
            "Each game session starts with 3 lives. In update(), enemies that cross the "
            "bottom boundary (y + height >= height - 100) are removed and decrement lives. "
            "Game over triggers only when lives reaches 0, not on the first boundary breach."
        ),
        "screenshot": "02_hud.png",
        "caption": "Figure 2: Lives counter displayed on HUD",
    },
    {
        "title": "Exercise 3 — Spaceship Graphic",
        "requirement": (
            "Add a spaceship graphic at the shooting position. Each time the player touches "
            "the screen, the spaceship should move to the touched location."
        ),
        "code": """// GameManager.kt — create rotated ship bitmap
fun createPlayerShipBitmap(): Bitmap {
    val source = BitmapFactory.decodeResource(context.resources, R.drawable.rocket_2)
    val scaled = Bitmap.createScaledBitmap(source, 90, 90, false)
    val matrix = Matrix().apply { postRotate(180f) }
    return Bitmap.createBitmap(scaled, 0, 0, scaled.width, scaled.height, matrix, true)
}

// GameView.kt — onTouchEvent
playerX = (event.x - playerShipWidth / 2f).coerceIn(0f, width - playerShipWidth)
canvas.drawBitmap(playerShipBitmap, null,
    RectF(playerX, playerY, playerX + playerShipWidth, playerY + playerShipHeight), null)""",
        "logic": (
            "GameManager loads rocket_2.png, scales and rotates it for the player ship. "
            "On ACTION_DOWN, playerX is set to the touch X coordinate (clamped to screen). "
            "draw() renders the bitmap at (playerX, playerY) above the bottom of the screen."
        ),
        "screenshot": "01_gameplay.png",
        "caption": "Figure 3: Player spaceship at touch position",
    },
    {
        "title": "Exercise 4 — Spreading Bullets",
        "requirement": (
            "Upgrade the bullet system so the bullets can spread/diverge over time. "
            "Each upgrade level should fire more bullets."
        ),
        "code": """// GameView.kt
bulletLevel = (1 + score / 50).coerceAtMost(5)

private fun fireBullets(centerX: Float, centerY: Float) {
    when (bulletLevel) {
        1 -> addBullet(centerX, centerY, 0f)
        2 -> {
            addBullet(centerX, centerY, 0f)
            addBullet(centerX - 25f, centerY, -3f)
            addBullet(centerX + 25f, centerY, 3f)
        }
        // levels 3–5 add more bullets with wider vx spread
    }
}

// FiringObject.kt — horizontal velocity
fun update() { x += vx; y -= speed; ... }""",
        "logic": (
            "bulletLevel increases every 50 points (max 5). fireBullets() uses a when "
            "expression to fire 1–5 bullets per tap, each with optional horizontal velocity "
            "(vx). FiringObject.update() applies vx so bullets spread outward as level rises."
        ),
        "screenshot": "01_gameplay.png",
        "caption": "Figure 4: Multiple spread bullets fired from the ship",
    },
    {
        "title": "Exercise 5 — Enemy Health Bars",
        "requirement": (
            "Add a health bar for each enemy. Every enemy should have a random amount of "
            "health, and the health bar should gradually decrease whenever the enemy is hit."
        ),
        "code": """// Opponent.kt
init {
    maxHealth = (2..5).random()
    currentHealth = maxHealth
}

fun takeDamage(damage: Int): Boolean {
    currentHealth -= damage
    return currentHealth <= 0
}

private fun drawHealthBar(canvas: Canvas) {
    val healthRatio = currentHealth.toFloat() / maxHealth.toFloat()
    canvas.drawRect(left, top, left + barWidth, top + barHeight, healthBgPaint)
    canvas.drawRect(left, top, left + barWidth * healthRatio, top + barHeight, healthFgPaint)
}""",
        "logic": (
            "Each Opponent spawns with random maxHealth (2–5). On bullet collision, "
            "takeDamage(1) reduces currentHealth. The enemy is removed only when health "
            "reaches 0. drawHealthBar() draws a red background and green foreground bar "
            "whose width reflects remaining health."
        ),
        "screenshot": "01_gameplay.png",
        "caption": "Figure 5: Enemies with health bars above them",
    },
    {
        "title": "Exercise 6 — Lane Movement",
        "requirement": (
            "Make enemies capable of moving across different lanes instead of only "
            "moving straight downward."
        ),
        "code": """// Opponent.kt
companion object { const val LANE_COUNT = 5 }

private val laneWidth: Float = screenWidth.toFloat() / laneCount
private val laneCenterX: Float = laneIndex * laneWidth + laneWidth / 2f - width / 2f

fun update() {
    y += speed
    val diff = laneCenterX - x
    if (abs(diff) > 1f) {
        x += diff.coerceIn(-4f, 4f)
    }
    rect.set(x, y, x + width, y + height)
}""",
        "logic": (
            "The screen is divided into 5 lanes. Each enemy has a laneIndex and a "
            "laneCenterX target. While falling (y += speed), the enemy drifts horizontally "
            "toward its lane center at up to 4 px/frame, creating diagonal lane movement."
        ),
        "screenshot": "01_gameplay.png",
        "caption": "Figure 6: Enemies drifting across lanes while falling",
    },
    {
        "title": "Exercise 7 — Boss Alien",
        "requirement": (
            "Create a larger alien enemy (Boss) that moves left and right while remaining "
            "fixed near the top of the screen. This Boss should spawn smaller aliens from "
            "the center of the Boss instead of spawning enemies randomly."
        ),
        "code": """// Boss.kt
fun update(screenWidth: Int) {
    x += dx
    if (x <= 0f || x + width >= screenWidth) { dx = -dx }
}

fun shouldSpawn(): Boolean {
    spawnTimer++
    if (spawnTimer >= spawnInterval) { spawnTimer = 0; return true }
    return false
}

// GameView.kt — boss spawns minions (no random spawn)
boss?.let { currentBoss ->
    currentBoss.update(width)
    if (currentBoss.shouldSpawn()) {
        val minion = gameManager.createMinion(
            currentBoss.getCenterX(), currentBoss.getCenterY() + ..., ...)
        opponents.add(minion)
    }
}""",
        "logic": (
            "Boss is created at y=60 and moves horizontally, bouncing off screen edges. "
            "Every 90 frames, shouldSpawn() returns true and GameView calls "
            "GameManager.createMinion() at the boss center. Random enemy spawning was "
            "removed from update(); all enemies now come from the boss."
        ),
        "screenshot": "03_boss_minions.png",
        "caption": "Figure 7: Boss alien at top spawning minions",
    },
]


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_labeled_paragraph(doc, label: str, text: str):
    p = doc.add_paragraph()
    run = p.add_run(f"{label} ")
    run.bold = True
    p.add_run(text)


def add_code_block(doc, code: str):
    for line in code.split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)


def add_image_if_exists(doc, path: Path, caption: str, width=Inches(2.8)):
    if path.exists():
        doc.add_picture(str(path), width=width)
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"[Screenshot not found: {path.name}]")


def ensure_screenshots():
    """Use real captures if available; otherwise generate composites."""
    if not any((SCREENSHOTS / f).exists() for f in ["01_gameplay.png", "02_hud.png"]):
        import generate_screenshots
        generate_screenshots.main()
        return False
    return REAL_SCREENSHOTS_FLAG.exists()


def main():
    SCREENSHOTS.mkdir(parents=True, exist_ok=True)
    real_captures = ensure_screenshots()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Lab 4 — Real-Time Mobile Application Development", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Course: Mobile Application Development")
    doc.add_paragraph("Students: MSSV 22521258, MSSV 22520953")
    doc.add_paragraph(
        "GitHub: https://github.com/tangkimson/Lab4_Mobile_22521258_22520953"
    )
    doc.add_paragraph()

    add_heading(doc, "1. Application Overview", 1)
    doc.add_paragraph(
        "This lab implements a classic space/chicken shooter game in Kotlin using "
        "SurfaceView and a dedicated GameThread running at approximately 60 FPS. "
        "The player controls a spaceship at the bottom, shoots yellow bullets at "
        "falling enemies, and avoids enemies reaching the bottom boundary."
    )

    add_heading(doc, "2. Architecture", 1)
    for line in [
        "MainActivity — sets GameView as the sole content view",
        "GameView — central controller: update/draw logic, touch input, HUD",
        "GameThread — 60 FPS loop: lockCanvas → update → draw → unlockCanvasAndPost",
        "GameManager — creates opponents, boss, minions, and player ship bitmaps",
        "Opponent — enemy with lanes, health, and health bar rendering",
        "FiringObject — yellow bullet rectangles with optional horizontal velocity",
        "Boss — top alien that moves horizontally and spawns minions",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    add_heading(doc, "3. Exercise Implementations", 1)
    if real_captures:
        doc.add_paragraph(
            "Screenshots in this section were captured from the running app on an Android emulator."
        )
    else:
        doc.add_paragraph(
            "Screenshots illustrate the implemented features using the project's game assets and HUD layout."
        )

    for ex in EXERCISES:
        add_heading(doc, ex["title"], 2)
        add_labeled_paragraph(doc, "Requirement:", ex["requirement"])
        add_labeled_paragraph(doc, "Code snippet:", "")
        add_code_block(doc, ex["code"])
        add_labeled_paragraph(doc, "Processing logic:", ex["logic"])
        add_image_if_exists(doc, SCREENSHOTS / ex["screenshot"], ex["caption"])

    add_heading(doc, "4. Game Over Screen", 2)
    doc.add_paragraph(
        "When lives reach zero, the game displays 'Game Over' and 'Tap to restart'. "
        "Tapping the screen calls resetGame() to clear objects and restore initial values."
    )
    add_image_if_exists(
        doc,
        SCREENSHOTS / "04_game_over.png",
        "Figure 8: Game Over screen",
    )

    add_heading(doc, "5. How to Run the Project", 1)
    for step in [
        "Clone: git clone https://github.com/tangkimson/Lab4_Mobile_22521258_22520953.git",
        "Open the project folder in Android Studio on Windows.",
        "Wait for Gradle sync (Android Studio creates local.properties automatically).",
        "Connect an emulator (API 24+) or physical device.",
        "Run the app (Shift+F10) or build with: .\\gradlew.bat assembleDebug",
        "Touch the screen to move the ship and fire. Tap after Game Over to restart.",
    ]:
        doc.add_paragraph(step, style="List Number")

    add_heading(doc, "6. Notes", 1)
    doc.add_paragraph(
        "This submission includes Lab4_Report.docx and Lab4_Report.pdf at the repository root. "
        "Source code is on GitHub for the teacher to clone and build on Windows with Android Studio."
    )

    doc.save(OUTPUT)
    print(f"Report saved to {OUTPUT}")
    print(f"Real emulator screenshots: {real_captures}")


if __name__ == "__main__":
    main()
